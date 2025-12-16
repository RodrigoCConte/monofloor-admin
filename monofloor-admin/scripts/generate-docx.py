#!/usr/bin/env python3
"""
Gerador de Relatório DOCX no formato Monofloor
Baseado no template oficial da empresa
"""

import sys
import json
import base64
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

def create_monofloor_report(data, output_path):
    """
    Cria relatório DOCX seguindo o template Monofloor
    """
    doc = Document()

    # Configurar fonte padrão Montserrat em todo o documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Montserrat'
    font.size = Pt(10)

    # =========================================================================
    # CABEÇALHO - Informações do Cliente
    # =========================================================================

    # Cliente
    p = doc.add_paragraph()
    run = p.add_run('Cliente: ')
    run.font.name = 'Montserrat'
    run.bold = True
    run = p.add_run(data.get('clientName', ''))
    run.font.name = 'Montserrat'

    # Endereço
    p = doc.add_paragraph()
    run = p.add_run('Endereço: ')
    run.font.name = 'Montserrat'
    run.bold = True
    run = p.add_run(data.get('address', ''))
    run.font.name = 'Montserrat'

    # Data da visita
    p = doc.add_paragraph()
    run = p.add_run('Data da visita: ')
    run.font.name = 'Montserrat'
    run.bold = True

    visit_date = data.get('visitDate', '')
    if visit_date:
        try:
            # Converter data ISO para formato brasileiro
            date_obj = datetime.fromisoformat(visit_date.replace('Z', '+00:00'))
            formatted_date = date_obj.strftime('%d/%m/%Y')
        except:
            formatted_date = visit_date
    else:
        formatted_date = datetime.now().strftime('%d/%m/%Y')

    run = p.add_run(formatted_date)
    run.font.name = 'Montserrat'

    # Espaço
    doc.add_paragraph()

    # =========================================================================
    # SAUDAÇÃO
    # =========================================================================

    p = doc.add_paragraph()
    run = p.add_run('Prezado(s) senhor(es),')
    run.font.name = 'Montserrat'
    run.font.size = Pt(12)
    run.bold = True

    # Espaço
    doc.add_paragraph()

    # =========================================================================
    # TEXTO INTRODUTÓRIO
    # =========================================================================

    intro_text = (
        "segue abaixo um detalhamento das observações feitas pelo nosso(a) colaborador(a) "
        "durante a visita à sua obra. As causas para a qualidade do acabamento ter ficado a "
        "desejar nessas regiões são diversas e serão listadas juntamente com a imagem "
        "correspondente a seguir."
    )

    p = doc.add_paragraph(intro_text)
    p.runs[0].font.name = 'Montserrat'

    reforco_text = (
        "Reforçamos a importância de considerar, além deste documento, as diretrizes do nosso "
        "contrato e manual de entrada, para que o substrato esteja devidamente preparado. "
        "Dessa forma, garantimos que a aplicação do revestimento ocorra com a qualidade e o "
        "acabamento esperados."
    )

    p = doc.add_paragraph(reforco_text)
    p.runs[0].font.name = 'Montserrat'

    # =========================================================================
    # DESCRIÇÃO DO SISTEMA
    # =========================================================================

    # Título
    p = doc.add_paragraph()
    run = p.add_run('Descrição do Sistema de Revestimento Monolítico')
    run.font.name = 'Montserrat'
    run.font.size = Pt(10)
    run.bold = True

    # Parágrafos descritivos
    desc_paragraphs = [
        (
            "O revestimento aplicado pela Monofloor segue um processo de execução cuidadosamente "
            "planejado para garantir a durabilidade e estética desejadas. A aplicação do sistema é "
            "realizada em camadas sequenciais, cada qual com sua função específica, passando por "
            "resinas agregantes, camadas de massa a base de resinas e componentes minerais que "
            "conferem a coloração e o aspecto artesanal do revestimento, cada uma delas seguida de "
            "lixamento manual, e camadas de selamento que incluem uma camada final de um verniz à "
            "base de poliuretano, que garante o acabamento, proteção e impermeabilidade do revestimento."
        ),
        (
            "O sistema completo tem uma espessura final que pode variar entre 1,5 e 3 milímetros a "
            "depender do tipo e qualidade do substrato em que está sendo aplicado. É necessário que "
            "elementos de piso como ralos, caixas elétricas, entre outros, sejam instalados rente ao "
            "substrato que receberá aplicação. Isso também vale para o encontro com outros revestimentos, "
            "qualquer degrau que haja irá permanecer, visto que o revestimento não pode ser utilizado "
            "para nenhum tipo de compensação de espessura."
        ),
        (
            "Importante lembrar que o revestimento é monolítico, ou seja, forma uma superfície contínua "
            "sem emendas, o que exige que, caso haja danos, a única forma de reparos ou reaplicação sem "
            "que haja marcas estéticas divergentes da aplicação original é com a reaplicação de toda a "
            "área contínua afetada."
        ),
        (
            "Durante o processo, é fundamental que o revestimento não entre em contato com qualquer tipo "
            "de contaminação, especialmente durante a aplicação das camadas finais de massa e selamento. "
            "No caso de contaminação por líquidos, a absorção dos mesmos por essas camadas pode causar "
            "manchas permanentes que comprometem a estética e a integridade do sistema como um todo. "
            "Contaminações em qualquer etapa não estão cobertas pela garantia e a resolução delas será "
            "cobrada como consta em contrato a depender da ocorrência."
        )
    ]

    for text in desc_paragraphs:
        p = doc.add_paragraph(text)
        p.runs[0].font.name = 'Montserrat'
        p.runs[0].font.size = Pt(10)

    # Espaço antes dos tópicos
    doc.add_paragraph()

    # =========================================================================
    # TÓPICOS COM IMAGENS E ANÁLISES
    # =========================================================================

    topic_groups = data.get('topicGroups', [])

    for topic_idx, topic_group in enumerate(topic_groups):
        # Título do tópico
        p = doc.add_paragraph()
        run = p.add_run(f"{topic_idx + 1}. {topic_group['topic']}")
        run.font.name = 'Montserrat'
        run.font.size = Pt(11)
        run.bold = True

        # Descrição do tópico (se houver)
        if topic_group.get('description'):
            p = doc.add_paragraph(topic_group['description'])
            p.runs[0].font.name = 'Montserrat'
            p.runs[0].font.size = Pt(10)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Frames do tópico
        frames = topic_group.get('frames', [])

        for frame in frames:
            # Adicionar imagem (agora usando base64 em vez de path)
            image_base64 = frame.get('imageBase64')
            if image_base64:
                try:
                    # Decodificar base64 e criar stream de bytes
                    image_data = base64.b64decode(image_base64)
                    image_stream = BytesIO(image_data)

                    # Adicionar imagem do stream
                    doc.add_picture(image_stream, width=Inches(5.5))
                    # Centralizar imagem
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                except Exception as e:
                    import traceback
                    print(f"Erro ao adicionar imagem: {e}")
                    print(f"Traceback completo:")
                    traceback.print_exc()
                    p = doc.add_paragraph('[Erro ao carregar imagem]')
                    p.runs[0].font.name = 'Montserrat'
                    p.runs[0].font.size = Pt(9)
                    p.runs[0].italic = True

            # Análise técnica
            analysis = frame.get('analysis', '')
            if analysis:
                p = doc.add_paragraph(analysis)
                p.runs[0].font.name = 'Montserrat'
                p.runs[0].font.size = Pt(10)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Espaço entre frames
            doc.add_paragraph()

        # Espaço maior entre tópicos
        if topic_idx < len(topic_groups) - 1:
            doc.add_paragraph()

    # =========================================================================
    # RODAPÉ / ENCERRAMENTO
    # =========================================================================

    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run('Atenciosamente,')
    run.font.name = 'Montserrat'
    run.font.size = Pt(11)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('Equipe Monofloor Revestimentos')
    run.font.name = 'Montserrat'
    run.font.size = Pt(11)
    run.bold = True

    # Salvar documento
    doc.save(output_path)
    print(f"✓ Relatório DOCX gerado: {output_path}")


def main():
    if len(sys.argv) != 3:
        print("Uso: python3 generate-docx.py <data.json> <output.docx>")
        sys.exit(1)

    data_path = sys.argv[1]
    output_path = sys.argv[2]

    # Ler dados
    with open(data_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Gerar relatório
    create_monofloor_report(data, output_path)


if __name__ == '__main__':
    main()
